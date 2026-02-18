# -*- coding: utf-8 -*-
"""
Генерация документации по аналитике клиентов и групп товаров
Создаёт красиво оформленный Word документ с аннотациями графиков
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Установка цвета фона ячейки"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def add_horizontal_line(paragraph):
    """Добавление горизонтальной линии"""
    p = paragraph._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '667eea')
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_documentation():
    """Создание документации"""
    doc = Document()

    # Настройка стилей
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Настройка полей страницы
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2)

    # ═══════════════════════════════════════════════════════════════
    # ТИТУЛЬНАЯ СТРАНИЦА
    # ═══════════════════════════════════════════════════════════════

    # Пустое пространство сверху
    for _ in range(6):
        doc.add_paragraph()

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('АНАЛИТИКА КЛИЕНТОВ\nИ ГРУПП ТОВАРОВ')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(102, 126, 234)  # #667eea

    doc.add_paragraph()

    # Подзаголовок
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Методология расчётов и описание визуализаций')
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_paragraph()
    doc.add_paragraph()

    # Информация о документе
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info.add_run('Период данных: 2023-2025\nКоличество графиков: 8\nИнтерактивный дашборд: customer_analytics_report.html')
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(120, 120, 120)

    # Разрыв страницы
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ОГЛАВЛЕНИЕ
    # ═══════════════════════════════════════════════════════════════

    toc_title = doc.add_paragraph()
    run = toc_title.add_run('СОДЕРЖАНИЕ')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(toc_title)

    doc.add_paragraph()

    contents = [
        ('1. Введение', '3'),
        ('2. Источники данных', '3'),
        ('3. График 1: Динамика статусов клиентов', '4'),
        ('4. График 2: RFM-кластеры по годам', '6'),
        ('5. График 3: Профили RFM-кластеров (Radar)', '8'),
        ('6. График 4: Миграция клиентов (Sankey)', '10'),
        ('7. График 5: Выручка по группам товаров', '12'),
        ('8. График 6: Сезонность групп товаров', '13'),
        ('9. График 7: Динамика выручки по годам', '14'),
        ('10. График 8: Сравнение новых и постоянных клиентов', '15'),
        ('11. Глоссарий терминов', '16'),
    ]

    for item, page in contents:
        p = doc.add_paragraph()
        p.add_run(item)
        p.add_run('.' * (60 - len(item)))
        p.add_run(page)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # 1. ВВЕДЕНИЕ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('1. ВВЕДЕНИЕ')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    intro_text = """Данный документ содержит подробное описание методологии расчётов и интерпретации визуализаций, представленных в интерактивном дашборде аналитики клиентов и групп товаров.

Цели анализа:
• Сегментация клиентов по степени лояльности и ценности (RFM-анализ)
• Отслеживание динамики поведения клиентов год от года
• Сравнение покупательского поведения новых и постоянных клиентов
• Анализ эффективности товарных групп

Период анализа: январь 2023 — декабрь 2025 (36 месяцев)
Количество клиентов: 1000
Количество товарных групп: 28"""

    doc.add_paragraph(intro_text)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════
    # 2. ИСТОЧНИКИ ДАННЫХ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('2. ИСТОЧНИКИ ДАННЫХ')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    # Таблица с данными
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'

    headers = ['Параметр', 'Значение']
    data = [
        ('Файл данных', 'test_sales_data.xlsx'),
        ('Структура', '3 года × 12 месяцев × 28 групп × 3 показателя'),
        ('Показатели', 'Количество в чеке, Сумма в чеке, Число чеков'),
        ('Идентификатор клиента', 'ID_клиента (CLIENT_0001 ... CLIENT_1000)'),
    ]

    # Заголовки
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '667eea')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Данные
    for row_idx, (param, value) in enumerate(data, 1):
        table.rows[row_idx].cells[0].text = param
        table.rows[row_idx].cells[1].text = value

    doc.add_paragraph()
    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГРАФИК 1: ДИНАМИКА СТАТУСОВ КЛИЕНТОВ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('3. ГРАФИК 1: Динамика статусов клиентов')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    # Описание
    h2 = doc.add_paragraph()
    run = h2.add_run('Описание')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""Stacked bar chart, показывающий распределение клиентов по статусам активности в каждом месяце. Позволяет отслеживать динамику притока новых клиентов, удержания постоянных и возврата ушедших.""")

    # Методология
    h2 = doc.add_paragraph()
    run = h2.add_run('Методология расчёта')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    # Таблица статусов
    table = doc.add_table(rows=5, cols=3)
    table.style = 'Table Grid'

    status_data = [
        ('Статус', 'Определение', 'Цвет'),
        ('Новый', 'Покупка в текущем месяце, НЕ было покупок 12 месяцев до этого', 'Зелёный (#2ecc71)'),
        ('Постоянный', 'Покупка в текущем месяце И была покупка за последние 6 месяцев', 'Фиолетовый (#9b59b6)'),
        ('Вернувшийся', 'Покупка в текущем месяце, НЕ было 6 месяцев, но был ранее', 'Синий (#3498db)'),
        ('Ушедший', 'Была покупка 6+ месяцев назад, НЕТ покупок последние 6 месяцев', 'Красный (#e74c3c)'),
    ]

    for row_idx, row_data in enumerate(status_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, '667eea')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # Формула
    h2 = doc.add_paragraph()
    run = h2.add_run('Формулы')
    run.bold = True
    run.font.size = Pt(13)

    formula_text = """
Для каждого клиента C в месяце M:

is_active_now = 1, если есть покупка в месяце M
active_last_6 = 1, если есть хотя бы одна покупка в месяцах [M-6, M-1]
active_last_12 = 1, если есть хотя бы одна покупка в месяцах [M-12, M-1]
active_before_6 = 1, если есть покупка ранее M-6

Статус определяется по алгоритму:
IF is_active_now = 1:
    IF active_last_12 = 0 → "Новый"
    ELIF active_last_6 = 1 → "Постоянный"
    ELSE → "Вернувшийся"
ELSE:
    IF active_before_6 = 1 AND active_last_6 = 0 → "Ушедший"
"""

    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    # Пример
    h2 = doc.add_paragraph()
    run = h2.add_run('Пример расчёта')
    run.bold = True
    run.font.size = Pt(13)

    example_text = """
Клиент CLIENT_0042, текущий месяц: Март 2024

История покупок:
• Январь 2023: была покупка ✓
• Февраль 2023 — Август 2023: нет покупок
• Сентябрь 2023: была покупка ✓
• Октябрь 2023 — Февраль 2024: нет покупок (6 месяцев)
• Март 2024: есть покупка ✓

Расчёт:
• is_active_now = 1 (есть покупка в марте 2024)
• active_last_6 = 0 (нет покупок с сентября 2023 по февраль 2024)
• active_last_12 = 1 (была покупка в сентябре 2023)
• active_before_6 = 1 (были покупки в 2023)

Результат: Статус = "Вернувшийся" (активен сейчас, не было 6 мес, но был ранее)
"""
    doc.add_paragraph(example_text)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГРАФИК 2: RFM-КЛАСТЕРЫ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('4. ГРАФИК 2: RFM-кластеры по годам')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    # Описание
    h2 = doc.add_paragraph()
    run = h2.add_run('Описание')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""Grouped bar chart, отображающий распределение клиентов по RFM-кластерам для каждого года. RFM-анализ — классический метод сегментации клиентов на основе их покупательского поведения.""")

    # Методология
    h2 = doc.add_paragraph()
    run = h2.add_run('Методология расчёта RFM')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    # Таблица RFM
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    rfm_data = [
        ('Показатель', 'Расшифровка', 'Формула расчёта'),
        ('R (Recency)', 'Давность последней покупки', 'Номер последнего месяца с покупкой (1-12)\n12 = декабрь = свежий клиент'),
        ('F (Frequency)', 'Частота покупок', 'Количество месяцев с покупками за год (1-12)'),
        ('M (Monetary)', 'Денежная ценность', 'Сумма всех покупок клиента за год (руб.)'),
    ]

    for row_idx, row_data in enumerate(rfm_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, '667eea')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # Кластеризация
    h2 = doc.add_paragraph()
    run = h2.add_run('Алгоритм кластеризации')
    run.bold = True
    run.font.size = Pt(13)

    clustering_text = """
1. Нормализация данных (StandardScaler):
   R_norm = (R - mean(R)) / std(R)
   F_norm = (F - mean(F)) / std(F)
   M_norm = (M - mean(M)) / std(M)

2. K-means кластеризация (k=5):
   Алгоритм минимизирует внутрикластерную дисперсию:
   argmin Σ Σ ||x - μᵢ||²

3. Присвоение названий кластерам:
   • Вычисляется средний Score = R_norm + F_norm + M_norm для каждого кластера
   • Кластеры ранжируются по Score (от большего к меньшему)
   • Присваиваются названия: Чемпионы → Лояльные → Перспективные →
     Требуют внимания → Потерянные
"""

    p = doc.add_paragraph()
    run = p.add_run(clustering_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()

    # Описание кластеров
    h2 = doc.add_paragraph()
    run = h2.add_run('Характеристики кластеров')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph()

    table = doc.add_table(rows=6, cols=4)
    table.style = 'Table Grid'

    cluster_data = [
        ('Кластер', 'R (Recency)', 'F (Frequency)', 'M (Monetary)'),
        ('Чемпионы', 'Высокий (недавно)', 'Высокий (часто)', 'Высокий (много)'),
        ('Лояльные', 'Средне-высокий', 'Высокий', 'Средне-высокий'),
        ('Перспективные', 'Высокий', 'Низкий', 'Средний'),
        ('Требуют внимания', 'Средний', 'Средний', 'Средний'),
        ('Потерянные', 'Низкий (давно)', 'Низкий (редко)', 'Низкий (мало)'),
    ]

    for row_idx, row_data in enumerate(cluster_data):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx].cells[col_idx]
            cell.text = cell_text
            if row_idx == 0:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, '667eea')
                cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    doc.add_paragraph()

    # Пример
    h2 = doc.add_paragraph()
    run = h2.add_run('Пример расчёта')
    run.bold = True
    run.font.size = Pt(13)

    example_text = """
Клиент CLIENT_0015, год 2024:

Исходные данные:
• Последний месяц покупки: Декабрь (12)
• Месяцев с покупками: 10 из 12
• Общая сумма: 2,450,000 руб.

RFM-показатели:
• R = 12 (максимум, свежий клиент)
• F = 10 (высокая частота)
• M = 2,450,000 (выше среднего)

После нормализации и кластеризации:
• R_norm = 1.2, F_norm = 0.9, M_norm = 1.1
• Score = 3.2 (высокий)

Результат: Кластер = "Чемпионы"
"""
    doc.add_paragraph(example_text)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГРАФИК 3: RADAR
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('5. ГРАФИК 3: Профили RFM-кластеров (Radar)')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Описание')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""Radar chart (лепестковая диаграмма), визуализирующий профиль каждого RFM-кластера по трём осям: Recency, Frequency, Monetary. Позволяет наглядно сравнить характеристики разных сегментов клиентов.""")

    h2 = doc.add_paragraph()
    run = h2.add_run('Методология расчёта')
    run.bold = True
    run.font.size = Pt(13)

    formula_text = """
1. Для каждого кластера вычисляются средние значения R, F, M:
   R_cluster = mean(R) для всех клиентов кластера
   F_cluster = mean(F) для всех клиентов кластера
   M_cluster = mean(M) для всех клиентов кластера

2. Нормализация в диапазон [0, 1] для сравнимости:
   R_norm = (R_cluster - min(R)) / (max(R) - min(R))
   F_norm = (F_cluster - min(F)) / (max(F) - min(F))
   M_norm = (M_cluster - min(M)) / (max(M) - min(M))

3. Построение radar chart с тремя осями и замкнутым контуром
"""

    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Интерпретация')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""• Чем больше площадь фигуры — тем ценнее сегмент клиентов
• "Чемпионы" имеют максимальную площадь (высокие R, F, M)
• "Потерянные" имеют минимальную площадь (низкие R, F, M)
• Форма фигуры показывает "перекосы" — например, высокий M при низком F означает редкие, но крупные покупки""")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГРАФИК 4: SANKEY
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('6. ГРАФИК 4: Миграция клиентов (Sankey)')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Описание')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""Sankey-диаграмма, показывающая потоки миграции клиентов между RFM-кластерами за три года (2023 → 2024 → 2025). Ширина потока пропорциональна количеству клиентов, совершивших переход.""")

    h2 = doc.add_paragraph()
    run = h2.add_run('Методология расчёта')
    run.bold = True
    run.font.size = Pt(13)

    formula_text = """
1. Создание таблицы миграции:
   migration[client_id, year] = cluster_name

2. Подсчёт переходов между кластерами:
   FOR each pair (cluster_from, cluster_to):
       count = COUNT(clients WHERE
                     cluster[year] = cluster_from AND
                     cluster[year+1] = cluster_to)

3. Формирование потоков Sankey:
   source[] = индексы исходных кластеров
   target[] = индексы целевых кластеров
   value[] = количество клиентов в переходе
"""

    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Пример расчёта')
    run.bold = True
    run.font.size = Pt(13)

    example_text = """
Переход 2023 → 2024:

Матрица переходов (пример):
                    2024
              Чемпионы  Лояльные  Перспективные  Требуют вним.  Потерянные
2023
Чемпионы          150       30            10              5           5
Лояльные           40      180            20             10          10
Перспективные      20       40           100             30          10
Требуют вним.       5       10            30            120          35
Потерянные          0        5            10             20         165

Интерпретация:
• 150 "Чемпионов" 2023 остались "Чемпионами" в 2024 (удержание)
• 30 "Чемпионов" 2023 стали "Лояльными" в 2024 (снижение)
• 40 "Лояльных" 2023 стали "Чемпионами" в 2024 (рост)
• 165 "Потерянных" остались "Потерянными" (критично!)
"""
    doc.add_paragraph(example_text)

    h2 = doc.add_paragraph()
    run = h2.add_run('Ключевые метрики миграции')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""• Retention Rate (удержание) = клиенты, оставшиеся в том же кластере / всего в кластере
• Upgrade Rate (рост) = клиенты, перешедшие в лучший кластер / всего в кластере
• Downgrade Rate (падение) = клиенты, перешедшие в худший кластер / всего в кластере
• Churn Rate (отток) = клиенты, ставшие "Потерянными" / всего клиентов""")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГРАФИК 5: ВЫРУЧКА ПО ГРУППАМ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('7. ГРАФИК 5: Выручка по группам товаров')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Описание')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""Bar chart, показывающий суммарную выручку по каждой из 28 товарных групп за весь период анализа. Позволяет выявить наиболее и наименее доходные категории.""")

    h2 = doc.add_paragraph()
    run = h2.add_run('Методология расчёта')
    run.bold = True
    run.font.size = Pt(13)

    formula_text = """
Выручка_группы = SUM(Сумма_в_чеке)
                 WHERE Группа = текущая_группа
                 GROUP BY Группа

Сортировка: по убыванию выручки

Дополнительные метрики (для анализа):
• Доля в выручке = Выручка_группы / Общая_выручка × 100%
• Средний чек группы = Выручка_группы / Число_чеков_группы
• Penetration = COUNT(DISTINCT клиентов с покупкой группы) / Всего_клиентов × 100%
"""

    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГРАФИК 6: СЕЗОННОСТЬ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('8. ГРАФИК 6: Сезонность групп товаров')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Описание')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""Heatmap (тепловая карта), показывающий распределение выручки каждой группы товаров по месяцам. Позволяет выявить сезонные паттерны продаж.""")

    h2 = doc.add_paragraph()
    run = h2.add_run('Методология расчёта')
    run.bold = True
    run.font.size = Pt(13)

    formula_text = """
1. Агрегация выручки по группам и месяцам:
   revenue[группа, месяц] = SUM(Сумма_в_чеке)
                           WHERE Группа = группа AND Месяц = месяц

2. Нормализация по строкам (для выявления паттернов):
   revenue_norm[группа, месяц] = revenue[группа, месяц] /
                                  SUM(revenue[группа, :]) × 100%

   Интерпретация: какой процент годовой выручки группы приходится на каждый месяц
   При равномерном распределении: ~8.3% на каждый месяц (100% / 12)

3. Цветовая шкала: YlOrRd (жёлтый → оранжевый → красный)
   • Жёлтый: низкая доля (< 7%)
   • Оранжевый: средняя доля (7-10%)
   • Красный: высокая доля (> 10%) — пик сезона
"""

    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГРАФИК 7: ДИНАМИКА ПО ГОДАМ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('9. ГРАФИК 7: Динамика выручки по годам')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Описание')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""Grouped bar chart, сравнивающий выручку по товарным группам в разрезе годов (2023, 2024, 2025). Позволяет оценить динамику роста/падения каждой категории.""")

    h2 = doc.add_paragraph()
    run = h2.add_run('Методология расчёта')
    run.bold = True
    run.font.size = Pt(13)

    formula_text = """
Выручка[группа, год] = SUM(Сумма_в_чеке)
                       WHERE Группа = группа AND Год = год

Дополнительные метрики:
• YoY Growth = (Выручка[год] - Выручка[год-1]) / Выручка[год-1] × 100%
• CAGR = (Выручка[2025] / Выручка[2023])^(1/2) - 1
"""

    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГРАФИК 8: НОВЫЕ VS ПОСТОЯННЫЕ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('10. ГРАФИК 8: Сравнение новых и постоянных клиентов')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Описание')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""Grouped bar chart, сравнивающий структуру покупок новых и постоянных клиентов по товарным группам. Показывает, какие категории предпочитают разные сегменты.""")

    h2 = doc.add_paragraph()
    run = h2.add_run('Методология расчёта')
    run.bold = True
    run.font.size = Pt(13)

    formula_text = """
1. Фильтрация данных по статусу:
   data_new = данные WHERE Статус = "Новый"
   data_regular = данные WHERE Статус = "Постоянный"

2. Агрегация чеков по группам:
   checks_new[группа] = SUM(Число_чеков) для data_new
   checks_regular[группа] = SUM(Число_чеков) для data_regular

3. Нормализация в проценты:
   pct_new[группа] = checks_new[группа] / SUM(checks_new) × 100%
   pct_regular[группа] = checks_regular[группа] / SUM(checks_regular) × 100%

4. Расчёт разницы (для выявления предпочтений):
   diff[группа] = pct_new[группа] - pct_regular[группа]

   • diff > 0: группу чаще покупают НОВЫЕ клиенты
   • diff < 0: группу чаще покупают ПОСТОЯННЫЕ клиенты
"""

    p = doc.add_paragraph()
    run = p.add_run(formula_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

    doc.add_paragraph()

    h2 = doc.add_paragraph()
    run = h2.add_run('Бизнес-интерпретация')
    run.bold = True
    run.font.size = Pt(13)

    doc.add_paragraph("""• Группы с высокой долей у новых клиентов — "точки входа", привлекающие новую аудиторию
• Группы с высокой долей у постоянных — "якорные категории", обеспечивающие лояльность
• Анализ помогает оптимизировать ассортимент и промо-активности для разных сегментов""")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════
    # ГЛОССАРИЙ
    # ═══════════════════════════════════════════════════════════════

    h1 = doc.add_paragraph()
    run = h1.add_run('11. ГЛОССАРИЙ ТЕРМИНОВ')
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(102, 126, 234)
    add_horizontal_line(h1)

    doc.add_paragraph()

    glossary = [
        ('RFM-анализ', 'Метод сегментации клиентов по трём параметрам: Recency (давность), Frequency (частота), Monetary (денежная ценность)'),
        ('Recency (R)', 'Время с момента последней покупки клиента'),
        ('Frequency (F)', 'Частота совершения покупок за период'),
        ('Monetary (M)', 'Общая сумма покупок клиента за период'),
        ('K-means', 'Алгоритм кластеризации, разбивающий данные на k групп с минимальной внутригрупповой дисперсией'),
        ('Sankey-диаграмма', 'Тип визуализации потоков, где ширина линии пропорциональна величине потока'),
        ('Heatmap', 'Тепловая карта — визуализация матрицы данных с помощью цветовой шкалы'),
        ('YoY (Year over Year)', 'Сравнение показателя с аналогичным периодом прошлого года'),
        ('Penetration', 'Доля клиентов, совершивших покупку в категории, от общего числа клиентов'),
        ('Retention Rate', 'Коэффициент удержания клиентов'),
        ('Churn Rate', 'Коэффициент оттока клиентов'),
        ('CAGR', 'Compound Annual Growth Rate — среднегодовой темп роста'),
    ]

    table = doc.add_table(rows=len(glossary)+1, cols=2)
    table.style = 'Table Grid'

    # Заголовок
    table.rows[0].cells[0].text = 'Термин'
    table.rows[0].cells[1].text = 'Определение'
    for cell in table.rows[0].cells:
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, '667eea')
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)

    # Данные
    for i, (term, definition) in enumerate(glossary, 1):
        table.rows[i].cells[0].text = term
        table.rows[i].cells[1].text = definition
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # Финальная строка
    final = doc.add_paragraph()
    final.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = final.add_run('— Конец документа —')
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.size = Pt(10)

    # Сохранение
    output_file = 'customer_analytics_documentation.docx'
    doc.save(output_file)
    print(f'Dokumentacia sohranena: {output_file}')
    return output_file

if __name__ == '__main__':
    create_documentation()
